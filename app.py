import json
from datetime import datetime
from io import BytesIO
from zoneinfo import ZoneInfo

import docx
import PyPDF2
import requests
import streamlit as st

try:
    import google.generativeai as genai
except ImportError:
    genai = None


APP_TITLE = "Agriculture Acarology Gujarati Article Writer"
DEFAULT_GEMINI_MODEL = "gemini-1.5-pro"
DEFAULT_PERPLEXITY_MODEL = "sonar-deep-research"


st.set_page_config(page_title=APP_TITLE, layout="wide")


def today_india():
    return datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d")


def ensure_state():
    defaults = {
        "topic_ideas": "",
        "topic_options": [],
        "selected_topic": "",
        "research_brief": "",
        "source_text": "",
        "draft_article": "",
        "editor_report": "",
        "final_article": "",
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)


def extract_text_from_pdf(uploaded_file):
    reader = PyPDF2.PdfReader(uploaded_file)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def extract_text_from_docx(uploaded_file):
    document = docx.Document(uploaded_file)
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def create_word_docx(article_text, title="Gujarati Newspaper Article"):
    document = docx.Document()
    document.add_heading(title, 0)
    document.add_paragraph(article_text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def call_gemini(api_key, model_name, prompt, temperature=0.45):
    if genai is None:
        raise RuntimeError("google-generativeai package is not installed.")
    if not api_key:
        raise RuntimeError("Gemini API key is missing.")
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(
        prompt,
        generation_config={"temperature": temperature},
    )
    return response.text.strip()


def call_perplexity(api_key, model_name, prompt):
    if not api_key:
        raise RuntimeError("Perplexity API key is missing.")

    response = requests.post(
        "https://api.perplexity.ai/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": model_name,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "You are a careful research assistant for an Indian Gujarati "
                        "newspaper agriculture column. Prefer current, source-grounded facts."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
        },
        timeout=90,
    )
    response.raise_for_status()
    data = response.json()
    return data["choices"][0]["message"]["content"].strip()


def parse_topic_options(topic_text):
    try:
        data = json.loads(topic_text)
        if isinstance(data, list):
            return [item for item in data if isinstance(item, dict) and item.get("title")]
        if isinstance(data, dict) and isinstance(data.get("topics"), list):
            return [item for item in data["topics"] if isinstance(item, dict) and item.get("title")]
    except json.JSONDecodeError:
        return []
    return []


def build_topic_prompt(date_text, region, crop_focus):
    return f"""
Today is {date_text} in India.

Search the web and find 6 timely article topics only related to agricultural acarology for farmers.

Region: {region}
Current season crop focus: {crop_focus}

Topic scope:
- Agricultural mites and acarology
- Crop mites, spider mites, eriophyid mites, predatory mites, mite outbreaks
- Mite management, farmer advisory, avoidable crop losses
- New technology, monitoring, biological control, IPM, weather-linked risk
- Current season crops, pest pressure, and farmer decision timing

Avoid topics that are only general insects, plant diseases, market prices, or policy unless mites/acarology are central.

Return strict JSON only. No markdown. No commentary.
Format:
[
  {{
    "title": "short English topic title",
    "gujarati_headline": "Gujarati headline idea",
    "why_now": "why this matters today/current season",
    "farmer_value": "how it helps farmers avoid losses",
    "technology_angle": "new technology, research, IPM, monitoring, biological control, or advisory angle",
    "verification_needed": "facts to verify before publication"
  }}
]
"""


def build_research_prompt(date_text, topic, region, source_text):
    return f"""
Today is {date_text} in India.

Deep research this selected agriculture acarology topic for a farmer-focused Gujarati newspaper article:
{topic}

Region: {region}

Extra source material from user:
{source_text[:12000]}

Prepare a research brief with:
- Current relevance for farmers and current season crops
- Mite/acarology background in simple language
- Crop damage symptoms and avoidable loss points
- Practical farmer advisory and IPM guidance
- Monitoring/scouting method
- New technology, biological control, predatory mites, digital tools, or research angle if relevant
- Key verified facts with source names or URLs
- Risks, cautions, uncertainty, and facts that need human verification
- Source names or URLs where available

Do not write the final article yet.
Do not recommend unsafe pesticide doses. If chemical control is discussed, keep it general and mention label/local expert guidance.
"""


def build_draft_prompt(topic, research_brief, source_text, word_count, style_rules):
    return f"""
You are an experienced Gujarati agriculture newspaper journalist writing in Agri Sandesh style.

Write a polished Gujarati newspaper article about this agriculture acarology topic:
{topic}

Target length: about {word_count} words.

Research brief:
{research_brief[:16000]}

Additional source material:
{source_text[:12000]}

Mandatory writing rules:
1. Entire output must be in fluent Gujarati.
2. Write in Agri Sandesh style for farmers and general readers.
3. Use a strong Gujarati headline.
4. Use flowing paragraphs only. Do not use bullet points, numbered lists, or dash lists.
5. Do not mention university, college, institute, or department names unless the editor explicitly asks.
6. Do not invent facts, figures, pesticide doses, scheme details, or quotes.
7. If a technical recommendation is uncertain, write it cautiously and advise local expert/label guidance.
8. Make it clearly related to agricultural acarology, mites, crop protection, and farmer loss prevention.
9. Explain symptoms, monitoring, prevention, IPM, and technology in farmer-friendly Gujarati.
10. Keep the language practical, clear, and suitable for publication.

Extra editor rules:
{style_rules}
"""


def build_editor_prompt(topic, article, research_brief):
    return f"""
You are a senior Gujarati newspaper editor and fact-checker.

Review this Gujarati agriculture article before publication.

Topic:
{topic}

Research brief:
{research_brief[:10000]}

Article:
{article[:16000]}

Give a concise editorial report in Gujarati with:
1. Fact risks
2. Weak or unclear parts
3. Tone/readability issues
4. Missing practical farmer value
5. Final rewrite instructions

Do not rewrite the article in this step.
"""


def build_final_rewrite_prompt(article, editor_report):
    return f"""
Rewrite the article below into a final Gujarati newspaper-ready version.

Current article:
{article[:16000]}

Editorial report and rewrite instructions:
{editor_report[:8000]}

Final rules:
1. Entire output must be Gujarati.
2. Keep an Agri Sandesh newspaper tone.
3. Use one strong headline.
4. Use continuous paragraphs only.
5. No bullet points, numbered lists, or dash lists.
6. Remove unsupported claims.
7. Improve clarity, flow, farmer usefulness, and acarology relevance.
"""


ensure_state()

st.title(APP_TITLE)
st.caption("Perplexity finds current acarology topics, Perplexity researches the selected topic, Gemini writes Gujarati Agri Sandesh-style article.")

with st.sidebar:
    st.header("API Keys")
    gemini_api_key = st.text_input("Gemini API Key", type="password")
    perplexity_api_key = st.text_input("Perplexity API Key", type="password")

    st.header("Models")
    gemini_model = st.text_input("Gemini model", DEFAULT_GEMINI_MODEL)
    perplexity_model = st.text_input("Perplexity model", DEFAULT_PERPLEXITY_MODEL)

    st.header("Article Settings")
    region = st.text_input("Region", "Gujarat, India")
    crop_focus = st.text_input("Current season crop focus", "cotton, chilli, vegetables, fruit crops, and other current season crops")
    word_count = st.slider("Approx word count", 450, 1300, 750, 50)
    style_rules = st.text_area(
        "Extra rules",
        "Keep it suitable for Agri Sandesh-style daily newspaper agriculture readers. Focus on mites, farmer advisory, avoidable losses, IPM, and practical field use.",
        height=90,
    )

tab_topic, tab_research, tab_write, tab_sources, tab_export = st.tabs(
    ["1. Perplexity Topic Search", "2. Perplexity Deep Research", "3. Gemini Article", "Optional Sources", "Export"]
)

with tab_topic:
    st.subheader("Step 1: Perplexity searches current acarology topics")
    st.write("This step searches only agriculture acarology topics useful for farmers, current crops, pest advisory, new technology, and avoidable crop losses.")

    if st.button("Search Acarology Topics With Perplexity", use_container_width=True):
        try:
            prompt = build_topic_prompt(today_india(), region, crop_focus)
            st.session_state.topic_ideas = call_perplexity(perplexity_api_key, perplexity_model, prompt)
            st.session_state.topic_options = parse_topic_options(st.session_state.topic_ideas)
        except Exception as exc:
            st.error(f"Perplexity topic search failed: {exc}")

    if st.session_state.topic_options:
        labels = [
            f"{item['title']} | {item.get('gujarati_headline', '')}"
            for item in st.session_state.topic_options
        ]
        selected_label = st.selectbox("Select one topic", labels)
        selected_index = labels.index(selected_label)
        selected_item = st.session_state.topic_options[selected_index]
        st.session_state.selected_topic = json.dumps(selected_item, ensure_ascii=False, indent=2)

        st.markdown("#### Selected Topic Details")
        st.info(
            f"Gujarati headline: {selected_item.get('gujarati_headline', '')}\n\n"
            f"Why now: {selected_item.get('why_now', '')}\n\n"
            f"Farmer value: {selected_item.get('farmer_value', '')}\n\n"
            f"Technology angle: {selected_item.get('technology_angle', '')}"
        )
    else:
        st.text_area("Raw topic search result", key="topic_ideas", height=260)
        st.text_area("Paste or type selected topic here if JSON parsing fails", key="selected_topic", height=110)

with tab_research:
    st.subheader("Step 2: Perplexity deep-researches your selected topic")
    st.write("After you select a topic, this step creates a source-grounded brief for farmers: symptoms, losses, scouting, IPM, technology, and cautions.")

    if st.button("Deep Research Selected Topic With Perplexity", use_container_width=True):
        if not st.session_state.selected_topic.strip():
            st.warning("Please select a topic first.")
        else:
            try:
                prompt = build_research_prompt(
                    today_india(),
                    st.session_state.selected_topic,
                    region,
                    st.session_state.source_text,
                )
                st.session_state.research_brief = call_perplexity(perplexity_api_key, perplexity_model, prompt)
            except Exception as exc:
                st.error(f"Perplexity research failed: {exc}")

    st.text_area("Perplexity deep research brief", key="research_brief", height=420)

with tab_write:
    st.subheader("Step 3: Gemini writes Gujarati Agri Sandesh-style article")

    col_a, col_b, col_c = st.columns(3)
    with col_a:
        if st.button("Generate Gujarati Article", use_container_width=True):
            if not st.session_state.research_brief.strip():
                st.warning("Please run Perplexity deep research first.")
            else:
                try:
                    prompt = build_draft_prompt(
                        st.session_state.selected_topic,
                        st.session_state.research_brief,
                        st.session_state.source_text,
                        word_count,
                        style_rules,
                    )
                    st.session_state.draft_article = call_gemini(gemini_api_key, gemini_model, prompt)
                except Exception as exc:
                    st.error(f"Draft generation failed: {exc}")

    with col_b:
        if st.button("Gujarati Editor Check", use_container_width=True):
            try:
                prompt = build_editor_prompt(
                    st.session_state.selected_topic,
                    st.session_state.draft_article,
                    st.session_state.research_brief,
                )
                st.session_state.editor_report = call_gemini(gemini_api_key, gemini_model, prompt, temperature=0.25)
            except Exception as exc:
                st.error(f"Editor check failed: {exc}")

    with col_c:
        if st.button("Final Gujarati Rewrite", use_container_width=True):
            try:
                prompt = build_final_rewrite_prompt(
                    st.session_state.draft_article,
                    st.session_state.editor_report,
                )
                st.session_state.final_article = call_gemini(gemini_api_key, gemini_model, prompt, temperature=0.35)
            except Exception as exc:
                st.error(f"Final rewrite failed: {exc}")

    st.markdown("#### Gemini Draft Article")
    st.text_area("Draft", key="draft_article", height=300)

    st.markdown("#### Editor Report")
    st.text_area("Report", key="editor_report", height=180)

    st.markdown("#### Final Article")
    st.text_area("Final", key="final_article", height=340)

with tab_sources:
    st.subheader("Optional: Add extra source material")
    pasted = st.text_area(
        "Paste research notes, article links, English text, Gujarati text, or official advisory text",
        height=240,
    )

    uploaded_files = st.file_uploader(
        "Upload PDF or DOCX source files",
        type=["pdf", "docx"],
        accept_multiple_files=True,
    )

    if st.button("Save Optional Sources", use_container_width=True):
        collected = [pasted.strip()] if pasted.strip() else []
        for uploaded_file in uploaded_files:
            try:
                if uploaded_file.name.lower().endswith(".pdf"):
                    collected.append(extract_text_from_pdf(uploaded_file))
                elif uploaded_file.name.lower().endswith(".docx"):
                    collected.append(extract_text_from_docx(uploaded_file))
            except Exception as exc:
                st.warning(f"Could not read {uploaded_file.name}: {exc}")
        st.session_state.source_text = "\n\n".join(text for text in collected if text.strip())
        st.success("Optional sources saved.")

    st.text_area("Saved source text", key="source_text", height=260)

with tab_export:
    st.subheader("Export")
    final_text = st.session_state.final_article.strip() or st.session_state.draft_article.strip()

    if final_text:
        st.download_button(
            "Download Final Article as Word",
            data=create_word_docx(final_text, "Gujarati Newspaper Article"),
            file_name="Gujarati_Newspaper_Article.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.markdown("#### Preview")
        st.info(final_text)
    else:
        st.warning("Generate a draft or final article first.")

    with st.expander("Saved workflow data as JSON"):
        st.code(
            json.dumps(
                {
                    "date": today_india(),
                    "region": region,
                    "selected_topic": st.session_state.selected_topic,
                    "research_brief": st.session_state.research_brief,
                    "draft_article": st.session_state.draft_article,
                    "editor_report": st.session_state.editor_report,
                    "final_article": st.session_state.final_article,
                },
                ensure_ascii=False,
                indent=2,
            ),
            language="json",
        )
